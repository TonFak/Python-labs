from docx import Document
from docx.shared import Inches

doc = Document('text.docx')

img_paragraph = doc.add_paragraph()
run = img_paragraph.add_run()
run.add_picture('Stud.jpg', width=Inches(1.25))

caption = doc.add_paragraph('Студсовет ОНК ИВТ')

doc.save('text.docx')